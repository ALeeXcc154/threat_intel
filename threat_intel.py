"""
威胁情报处理器 - 修复版
1. 还原了原始解析逻辑 (支持 ** 小标题识别)
2. 增加了跨 URL 相同标题合并逻辑
3. 还原了原版 Word 精确排版格式
"""

import re
import json
import requests
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn


import socket
from urllib.parse import urlparse
import ipaddress

class ThreatIntelProcessor:
    """威胁情报处理器"""
    
    def __init__(self):
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
    
    def _validate_url_security(self, url: str):
        """[SSRF防护] 校验URL安全性，禁止私有IP访问"""
        try:
            parsed = urlparse(url)
            if parsed.scheme not in ('http', 'https'):
                raise ValueError("仅支持 HTTP/HTTPS 协议")
            
            hostname = parsed.hostname
            if not hostname:
                raise ValueError("URL 格式无效")
                
            # 解析 IP 地址
            try:
                ip = socket.gethostbyname(hostname)
            except socket.gaierror:
                raise ValueError(f"无法解析域名: {hostname}")
                
            # 检查是否为私有 IP
            ip_obj = ipaddress.ip_address(ip)
            if ip_obj.is_private or ip_obj.is_loopback or ip_obj.is_link_local:
                raise ValueError(f"安全拦截: 禁止访问私有地址 ({hostname} -> {ip})")
                
        except Exception as e:
            raise ValueError(f"URL 安全校验失败: {str(e)}")

    def fetch_and_extract_data(self, url: str) -> Tuple[Optional[str], List[str]]:
        """从URL获取并提取原始 richText 数据"""
        messages = [f"正在请求URL: {url}"]
        try:
            # 执行 SSRF 安全校验
            self._validate_url_security(url)
            
            response = requests.get(url, headers=self.headers, timeout=30)
            response.raise_for_status()
            messages.append("网页内容获取成功！")
            
            match = re.search(r'window\.__INITIAL_STATE__\s*=\s*({.*?});?</script>', response.text, re.DOTALL)
            
            if match:
                messages.append("成功匹配到 __INITIAL_STATE__ 数据！")
                json_str = match.group(1)
                data = json.loads(json_str)
                rich_text = data["data"]["resData"]["data"]["threatDetail"]["richText"]
                messages.append("成功提取 'richText' 内容！")
                return rich_text, messages
            else:
                messages.append("错误：在页面中未找到 '__INITIAL_STATE__' 数据。")
                return None, messages
        except Exception as e:
            messages.append(f"请求失败: {e}")
            return None, messages

    def parse_content_to_entries(self, rich_text_content: str, url: str) -> List[Dict]:
        """解析内容并提取条目 (还原原版解析算法)"""
        lines = rich_text_content.split('\n')
        
        current_section = None
        main_title = "威胁情报"
        sub_title = None
        content_lines = []
        reference_link = None
        entries = []
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
                
            # 1. 识别大标题 (识别 ===== 及其上一行)
            if "=====" in line and (i + 1 < len(lines) and lines[i+1].strip() == ""):
                if sub_title or content_lines:
                    entries.append({
                        "main_title": main_title,
                        "sub_title": sub_title,
                        "content_lines": list(content_lines),
                        "reference_link": reference_link,
                        "source_url": url
                    })
                    sub_title, content_lines, reference_link = None, [], None
                
                if i > 0:
                    main_title = lines[i-1].strip()
                i += 1
                continue
                
            # 2. 识别小标题 (以 ** 开头，且排除特定关键词)
            if (line and main_title is not None and 
                (line.startswith("**") or (i + 1 < len(lines) and lines[i+1].strip().startswith("**"))) and
                not line.startswith("事件概述:") and 
                not line.startswith("参考链接:")):
                
                if sub_title or content_lines:
                    entries.append({
                        "main_title": main_title,
                        "sub_title": sub_title,
                        "content_lines": list(content_lines),
                        "reference_link": reference_link,
                        "source_url": url
                    })
                    sub_title, content_lines, reference_link = None, [], None
                
                if line.startswith("**"):
                    sub_title = line
                else:
                    sub_title = line
                    if i + 1 < len(lines) and lines[i+1].strip().startswith("**"):
                        sub_title += "\n" + lines[i+1].strip()
                        i += 1
                i += 1
                continue
                
            # 3. 识别正文开始
            if line.startswith("事件概述:"):
                current_section = "content"
                content_lines.append(line)
                i += 1
                continue
                
            # 4. 收集正文或参考链接
            if current_section == "content":
                if line.startswith("参考链接:"):
                    reference_link = line.replace("参考链接:", "").strip()
                    current_section = "reference"
                else:
                    content_lines.append(line)
                i += 1
                continue
                
            i += 1
            
        # 添加最后一个条目
        if sub_title or content_lines:
            entries.append({
                "main_title": main_title,
                "sub_title": sub_title,
                "content_lines": content_lines,
                "reference_link": reference_link,
                "source_url": url
            })
        
        return entries

    def merge_entries_by_main_title(self, entries: List[Dict]) -> List[Dict]:
        """合并相同大标题下的条目 (跨 URL 合并)"""
        merged_dict = {}
        
        for entry in entries:
            mt = entry["main_title"]
            st = entry["sub_title"]
            
            if mt not in merged_dict:
                merged_dict[mt] = {"main_title": mt, "sub_titles": {}}
            
            if st not in merged_dict[mt]["sub_titles"]:
                merged_dict[mt]["sub_titles"][st] = {
                    "sub_title": st,
                    "content_lines": [],
                    "reference_links": []
                }
            
            if entry["content_lines"]:
                merged_dict[mt]["sub_titles"][st]["content_lines"].extend(entry["content_lines"])
            if entry["reference_link"]:
                merged_dict[mt]["sub_titles"][st]["reference_links"].append(entry["reference_link"])
        
        # 转换为列表格式供 Word 生成
        result = []
        for mt_data in merged_dict.values():
            for st_data in mt_data["sub_titles"].values():
                result.append({
                    "main_title": mt_data["main_title"],
                    "sub_title": st_data["sub_title"],
                    "content_lines": st_data["content_lines"],
                    "reference_links": st_data["reference_links"]
                })
        return result

    def add_table_of_contents(self, doc):
        """添加目录域代码 (Office 自动更新)"""
        try:
            toc_title = doc.add_paragraph()
            run = toc_title.add_run("目  录")
            run.font.size, run.font.bold = Pt(16), True
            run.font.name = u'微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
            toc_title.alignment = 1
            
            toc_note = doc.add_paragraph("（目录标级已设置，正文需手动配置）")
            run_note = toc_note.runs[0]
            run_note.font.size, run_note.font.italic = Pt(10), True
            run_note.font.color.rgb = RGBColor(128, 128, 128)
            toc_note.alignment = 1
            
            doc.add_paragraph()
            from docx.oxml.shared import OxmlElement
            
            toc_para = doc.add_paragraph()
            fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText'); instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
            fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
            fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
            
            r = toc_para.add_run()
            r._r.append(fldChar1); r._r.append(instrText); r._r.append(fldChar2); r._r.append(fldChar3)
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        except Exception as e:
            print(f"添加目录失败: {e}")

    def generate_document(self, urls: List[str], output_path: str, progress_callback=None) -> Tuple[bool, List[str]]:
        """主入口：处理 URL 列表并生成 Word 文档"""
        messages = []
        all_raw_entries = []
        
        try:
            # 1. 抓取与解析阶段
            for i, url in enumerate(urls):
                if progress_callback:
                    progress_callback(int((i / len(urls)) * 40), f"解析中: {url}")
                
                rich_text, msg_list = self.fetch_and_extract_data(url)
                messages.extend(msg_list)
                
                if rich_text:
                    raw_entries = self.parse_content_to_entries(rich_text, url)
                    all_raw_entries.extend(raw_entries)

            if not all_raw_entries:
                return False, ["未获取到任何有效内容"]

            # 2. 合并阶段
            if progress_callback:
                progress_callback(50, "正在合并相同大标题的内容...")
            merged_entries = self.merge_entries_by_main_title(all_raw_entries)

            # 3. 生成 Word 文档
            if progress_callback:
                progress_callback(70, "正在渲染 Word 排版...")
                
            doc = Document()
            # 设置默认字体
            doc.styles['Normal'].font.name = u'微软雅黑'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
            doc.styles['Normal'].font.size = Pt(12)
            doc.styles['Normal'].paragraph_format.line_spacing = 1.5
            
            # 添加目录
            self.add_table_of_contents(doc)
            
            current_main_title = None
            main_title_first_entry = {} # 记录大标题下的首个条目

            for idx, entry in enumerate(merged_entries):
                # 如果大标题变化
                if entry["main_title"] and entry["main_title"] != current_main_title:
                    current_main_title = entry["main_title"]
                    # 除了第一个大标题，其他大标题前都分页
                    if idx > 0:
                        doc.add_page_break()
                    
                    # 添加红色大标题
                    title_para = doc.add_paragraph(style='Heading 1')
                    title_run = title_para.add_run(f"【{current_main_title}】")
                    title_run.font.size, title_run.font.bold = Pt(18), True
                    title_run.font.color.rgb = RGBColor(255, 0, 0)
                    title_run.font.name = u'微软雅黑'
                    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
                    
                    main_title_first_entry[current_main_title] = True
                
                # 添加小标题
                if entry["sub_title"]:
                    # 排版逻辑：同一大标题下的后续小标题需要另起一页
                    if not main_title_first_entry[current_main_title]:
                        doc.add_page_break()
                    main_title_first_entry[current_main_title] = False
                    
                    subtitle_para = doc.add_paragraph(style='Heading 2')
                    clean_st = entry["sub_title"].replace("**", "")
                    sub_run = subtitle_para.add_run(clean_st)
                    sub_run.font.size, sub_run.font.bold = Pt(14), True
                    sub_run.font.name = u'微软雅黑'
                    sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

                # 正文
                if entry["content_lines"]:
                    doc.add_paragraph() # 空行
                    content_text = '\n'.join(entry["content_lines"])
                    p = doc.add_paragraph(content_text)
                    p.paragraph_format.first_line_indent = Pt(24)
                    p.paragraph_format.line_spacing = 1.5

                # 参考链接
                if entry["reference_links"]:
                    doc.add_paragraph()
                    links = "; ".join(filter(None, entry["reference_links"]))
                    ref_text = "参考链接: " + links
                    ref_p = doc.add_paragraph()
                    ref_run = ref_p.add_run(ref_text)
                    ref_run.font.size, ref_run.font.italic = Pt(10), True
                    ref_run.font.color.rgb = RGBColor(0, 0, 0)

            doc.save(output_path)
            messages.append(f"文档已保存至: {output_path}")
            if progress_callback:
                progress_callback(100, "处理完成！")
            return True, messages

        except Exception as e:
            messages.append(f"生成失败: {str(e)}")
            return False, messages

    def process_url(self, url: str) -> tuple[bool, List[Dict], str]:
        """处理单个URL (用于预览或测试)"""
        try:
            rich_text, _ = self.fetch_and_extract_data(url)
            if not rich_text:
                return False, [], "数据提取为空"
            entries = self.parse_content_to_entries(rich_text, url)
            return True, entries, ""
        except Exception as e:
            return False, [], str(e)
